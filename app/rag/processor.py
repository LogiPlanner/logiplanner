import os
import re
import ipaddress
import socket
from urllib.parse import urlparse
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional, Tuple

import httpx
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LCDocument

from app.core.config import settings


# ─── Supported file extensions → loader mapping ───
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "docx",
    ".txt": "txt",
    ".md": "markdown",
}

# Maximum file size: 20MB
MAX_FILE_SIZE = 20 * 1024 * 1024

# GitHub repo processing constants
MAX_REPO_FILES = 100  # Limit total files to prevent huge repos
MAX_FILE_SIZE_REPO = 200_000  # ~200KB per file
ALLOWED_EXTENSIONS = {".py", ".md", ".txt", ".json", ".yaml", ".yml"}
IGNORE_DIRS = {".git", "__pycache__", "node_modules", "venv", ".env", "dist", "build"}


def get_doc_type(filename: str) -> Optional[str]:
    """Get document type from filename extension."""
    ext = os.path.splitext(filename)[1].lower()
    return SUPPORTED_EXTENSIONS.get(ext)


def validate_file(filename: str, file_size: int) -> tuple[bool, str]:
    """
    Validate a file before processing.
    Returns (is_valid, error_message).
    """
    doc_type = get_doc_type(filename)
    if not doc_type:
        supported = ", ".join(SUPPORTED_EXTENSIONS.keys())
        return False, f"Unsupported file type. Supported: {supported}"

    if file_size > MAX_FILE_SIZE:
        return False, f"File too large. Maximum size: {MAX_FILE_SIZE // (1024*1024)}MB"

    if file_size == 0:
        return False, "File is empty"

    return True, ""


def _load_docx(file_path: str) -> List[LCDocument]:
    """Load a DOCX file using python-docx for robust handling of complex documents."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    full_text = "\n\n".join(paragraphs)
    if not full_text.strip():
        raise ValueError("No readable text found in DOCX file")
    return [LCDocument(page_content=full_text, metadata={"source": file_path})]


def load_document(file_path: str, filename: str) -> List[LCDocument]:
    """
    Load a document using the appropriate LangChain loader.
    Returns a list of LangChain Document objects (one per page for PDFs).
    """
    doc_type = get_doc_type(filename)

    if doc_type == "pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif doc_type == "docx":
        return _load_docx(file_path)
    elif doc_type in ("txt", "markdown"):
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    else:
        raise ValueError(f"Unsupported document type: {doc_type}")


def split_documents(
    documents: List[LCDocument],
    chunk_size: int = None,
    chunk_overlap: int = None,
) -> List[LCDocument]:
    """
    Split documents into chunks using RecursiveCharacterTextSplitter.
    This preserves semantic coherence by splitting on natural boundaries
    (paragraphs → sentences → words) in that order.
    """
    chunk_size = chunk_size or settings.RAG_CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.RAG_CHUNK_OVERLAP

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
        is_separator_regex=False,
    )

    return splitter.split_documents(documents)


def enrich_metadata(
    chunks: List[LCDocument],
    team_id: int,
    document_id: int,
    filename: str,
    uploader_email: str,
    doc_type: str,
    doc_summary: str = "",
) -> List[LCDocument]:
    """
    Enrich each chunk with metadata for filtering and traceability.

    Metadata fields added to each chunk:
    - team_id: Which team this document belongs to
    - document_id: Database ID of the parent document record
    - filename: Original filename for source citations
    - uploader_email: Who uploaded this document
    - doc_type: pdf/docx/txt/markdown
    - chunk_index: Position of this chunk in the document
    - page_number: Page number (for PDFs, from loader metadata)
    - uploaded_at: ISO timestamp of when the document was ingested
    - doc_summary: One-sentence summary of the full document

    When RAG_CONTEXTUAL_HEADERS is enabled, the chunk text is prepended with the
    document title and summary so the embedding captures document-level context.
    """
    uploaded_at = datetime.now(timezone.utc).isoformat()

    def _build_contextual_header(filename_value: str, summary_value: str) -> str:
        header_parts = [f"[Document: {filename_value}]"]
        if summary_value:
            header_parts.append(f"[Summary: {summary_value}]")
        return " ".join(header_parts)

    for i, chunk in enumerate(chunks):
        # Preserve any existing metadata from loaders (e.g., page number from PDF)
        existing_meta = chunk.metadata or {}

        chunk.metadata = {
            "team_id": team_id,
            "document_id": document_id,
            "filename": filename,
            "uploader_email": uploader_email,
            "doc_type": doc_type,
            "chunk_index": i,
            "page_number": existing_meta.get("page", 0),
            "uploaded_at": uploaded_at,
            # Keep source from loader if present
            "source": existing_meta.get("source", filename),
            "doc_summary": doc_summary,
        }

        # Contextual chunk headers: prepend document title + summary to the
        # chunk text so the embedding vector captures document-level context.
        if settings.RAG_CONTEXTUAL_HEADERS:
            header = _build_contextual_header(filename, doc_summary)
            chunk.page_content = f"{header}\n\n{chunk.page_content}"

    return chunks


def apply_document_summary(chunks: List[LCDocument], doc_summary: str) -> List[LCDocument]:
    """Update chunk metadata and contextual headers after a summary is generated."""
    normalized_summary = (doc_summary or "").strip()

    for chunk in chunks:
        if chunk.metadata is None:
            chunk.metadata = {}

        filename = chunk.metadata.get("filename") or chunk.metadata.get("source") or "Document"
        chunk.metadata["doc_summary"] = normalized_summary

        if not settings.RAG_CONTEXTUAL_HEADERS:
            continue

        body = chunk.page_content or ""
        if body.startswith("[Document: "):
            _, separator, remainder = body.partition("\n\n")
            if separator:
                body = remainder

        header_parts = [f"[Document: {filename}]"]
        if normalized_summary:
            header_parts.append(f"[Summary: {normalized_summary}]")
        header = " ".join(header_parts)
        chunk.page_content = f"{header}\n\n{body}"

    return chunks


def process_document(
    file_path: str,
    filename: str,
    team_id: int,
    document_id: int,
    uploader_email: str,
) -> List[LCDocument]:
    """
    Full document processing pipeline:
    1. Load the document
    2. Split into chunks
    3. Enrich with metadata

    Returns list of processed, metadata-enriched chunks ready for embedding.
    """
    doc_type = get_doc_type(filename)
    if not doc_type:
        raise ValueError(f"Unsupported file: {filename}")

    # Step 1: Load
    raw_docs = load_document(file_path, filename)

    if not raw_docs:
        raise ValueError(f"No content could be extracted from {filename}")

    # Step 2: Split
    chunks = split_documents(raw_docs)

    if not chunks:
        raise ValueError(f"Document {filename} produced no chunks after splitting")

    # Step 3: Enrich metadata
    enriched = enrich_metadata(
        chunks=chunks,
        team_id=team_id,
        document_id=document_id,
        filename=filename,
        uploader_email=uploader_email,
        doc_type=doc_type,
    )

    return enriched


def process_text(
    text: str,
    title: str,
    team_id: int,
    document_id: int,
    uploader_email: str,
) -> List[LCDocument]:
    """
    Process raw text (e.g., notes, context) into chunks.
    Used for ingesting text directly without a file.
    """
    if not text or not text.strip():
        raise ValueError("Text content is empty")

    # Create a single LangChain document from the text
    doc = LCDocument(
        page_content=text,
        metadata={"source": title},
    )

    # Split
    chunks = split_documents([doc])

    # Enrich
    enriched = enrich_metadata(
        chunks=chunks,
        team_id=team_id,
        document_id=document_id,
        filename=title,
        uploader_email=uploader_email,
        doc_type="text",
    )

    return enriched


# ─── Google Drive URL Parsing & Download ───

# Patterns to extract file IDs from various Google Drive/Docs URLs
_DRIVE_PATTERNS = [
    # Google Docs: /document/d/{ID}/...
    re.compile(r"docs\.google\.com/document/d/([a-zA-Z0-9_-]+)"),
    # Google Sheets: /spreadsheets/d/{ID}/...
    re.compile(r"docs\.google\.com/spreadsheets/d/([a-zA-Z0-9_-]+)"),
    # Google Slides: /presentation/d/{ID}/...
    re.compile(r"docs\.google\.com/presentation/d/([a-zA-Z0-9_-]+)"),
    # Google Drive file: /file/d/{ID}/...
    re.compile(r"drive\.google\.com/file/d/([a-zA-Z0-9_-]+)"),
    # Google Drive folder: /drive/folders/{ID}
    re.compile(r"drive\.google\.com/drive/folders/([a-zA-Z0-9_-]+)"),
    # Google Drive open: ?id={ID}
    re.compile(r"drive\.google\.com/open\?id=([a-zA-Z0-9_-]+)"),
    # Google Drive uc: ?id={ID}
    re.compile(r"drive\.google\.com/uc\?.*?id=([a-zA-Z0-9_-]+)"),
]


def parse_drive_url(url: str) -> Tuple[str, str]:
    """
    Parse a Google Drive/Docs URL and return (file_id, url_type).

    url_type is one of: 'document', 'spreadsheet', 'presentation', 'file', 'folder'

    Raises ValueError if the URL is not a recognized Google Drive URL.
    """
    url = url.strip()

    # Determine the type from the URL
    if "docs.google.com/document" in url:
        url_type = "document"
    elif "docs.google.com/spreadsheets" in url:
        url_type = "spreadsheet"
    elif "docs.google.com/presentation" in url:
        url_type = "presentation"
    elif "drive.google.com/drive/folders" in url:
        url_type = "folder"
    elif "drive.google.com" in url:
        url_type = "file"
    else:
        raise ValueError(
            "Not a recognized Google Drive URL. "
            "Please use a link from drive.google.com or docs.google.com."
        )

    # Extract file ID
    for pattern in _DRIVE_PATTERNS:
        match = pattern.search(url)
        if match:
            return match.group(1), url_type

    raise ValueError(
        "Could not extract a file ID from this URL. "
        "Please use a standard Google Drive share link."
    )


def list_folder_files(folder_id: str) -> List[Tuple[str, str, str]]:
    """
    List files in a public Google Drive folder.

    Returns a list of (file_id, filename, mime_hint) tuples.
    Uses the public folder HTML page to scrape file IDs.
    """
    folder_url = f"https://drive.google.com/drive/folders/{folder_id}"

    try:
        with httpx.Client(follow_redirects=True, timeout=30.0) as client:
            resp = client.get(folder_url)
    except httpx.RequestError as e:
        raise ValueError(f"Failed to access Drive folder: {e}")

    if resp.status_code in (401, 403):
        raise ValueError(
            "Access denied. Make sure the folder is shared publicly: "
            "Folder → Share → General access → Anyone with the link."
        )
    if resp.status_code == 404:
        raise ValueError("Folder not found. Check the link and try again.")
    if resp.status_code != 200:
        raise ValueError(f"Google Drive returned HTTP {resp.status_code}.")

    html = resp.text

    # Google Drive folder pages embed file IDs in various attributes.
    # Look for /file/d/{ID} patterns and data-id attributes in the HTML.
    file_ids_seen = set()
    results = []

    # Pattern 1: /file/d/{ID}/ links in the HTML
    for m in re.finditer(r'/file/d/([a-zA-Z0-9_-]{10,})', html):
        fid = m.group(1)
        if fid not in file_ids_seen:
            file_ids_seen.add(fid)
            results.append((fid, f"file_{fid[:8]}", "file"))

    # Pattern 2: data-id="{ID}" attributes (Google Drive uses these)
    for m in re.finditer(r'data-id="([a-zA-Z0-9_-]{10,})"', html):
        fid = m.group(1)
        if fid not in file_ids_seen and fid != folder_id:
            file_ids_seen.add(fid)
            results.append((fid, f"file_{fid[:8]}", "file"))

    if not results:
        # If HTML scraping found nothing, the folder may be empty or
        # Google served a JS-only page. Give a clear error.
        if "sign in" in html[:3000].lower() or "accounts.google.com" in html[:3000]:
            raise ValueError(
                "This folder is not publicly shared. "
                "Set sharing to 'Anyone with the link' in Google Drive."
            )
        raise ValueError(
            "No files found in this folder. The folder may be empty, "
            "or Google served a page we couldn't parse. "
            "Try sharing individual file links instead."
        )

    return results


def _build_download_url(file_id: str, url_type: str) -> str:
    """Build the download/export URL based on the file type."""
    if url_type == "document":
        return f"https://docs.google.com/document/d/{file_id}/export?format=txt"
    elif url_type == "spreadsheet":
        return f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=csv"
    elif url_type == "presentation":
        return f"https://docs.google.com/presentation/d/{file_id}/export/txt"
    else:
        # Regular Drive file — direct download
        return f"https://drive.google.com/uc?export=download&id={file_id}"


def process_drive_url(
    url: str,
    team_id: int,
    document_id: int,
    uploader_email: str,
) -> Tuple[List[LCDocument], str, str, int]:
    """
    Download a public Google Drive document and process it into chunks.

    Returns: (chunks, filename, doc_type, file_size_bytes)

    Raises ValueError on download failure or private file.
    """
    file_id, url_type = parse_drive_url(url)
    download_url = _build_download_url(file_id, url_type)

    # For generic "file" type (e.g. files from folder scraping), the /uc endpoint
    # returns 500 for Google Workspace files (Docs, Sheets, Slides).
    # Build a list of URLs to try in order.
    urls_to_try = [download_url]
    if url_type == "file":
        urls_to_try.extend([
            f"https://docs.google.com/document/d/{file_id}/export?format=txt",
            f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=csv",
            f"https://docs.google.com/presentation/d/{file_id}/export/txt",
        ])

    resp = None
    used_url_type = url_type
    try:
        with httpx.Client(follow_redirects=True, timeout=60.0) as client:
            for i, try_url in enumerate(urls_to_try):
                resp = client.get(try_url)
                if resp.status_code == 200:
                    # Track which export format succeeded for naming
                    if i == 1:
                        used_url_type = "document"
                    elif i == 2:
                        used_url_type = "spreadsheet"
                    elif i == 3:
                        used_url_type = "presentation"
                    break
                # Only retry on 500 (server error) or 404 from export endpoints
                if resp.status_code not in (500, 404) or i == len(urls_to_try) - 1:
                    break
    except httpx.TimeoutException:
        raise ValueError("Download timed out. The file may be too large or the server is slow.")
    except httpx.RequestError as e:
        raise ValueError(f"Failed to connect to Google Drive: {e}")

    if resp.status_code == 404:
        raise ValueError(
            "File not found. Make sure the Google Drive link is correct "
            "and the file hasn't been deleted."
        )
    if resp.status_code in (401, 403):
        raise ValueError(
            "Access denied. Make sure the file is shared publicly: "
            "File → Share → General access → Anyone with the link."
        )
    if resp.status_code != 200:
        raise ValueError(f"Google Drive returned HTTP {resp.status_code}.")

    # ── Enforce 100 MB file size limit ──
    MAX_DRIVE_BYTES = 100 * 1024 * 1024  # 100 MB
    content_length = resp.headers.get("content-length")
    if content_length and int(content_length) > MAX_DRIVE_BYTES:
        raise ValueError(
            f"File is too large ({int(content_length) // (1024*1024)} MB). "
            "Maximum allowed size is 100 MB."
        )

    content = resp.text
    content_bytes = len(content.encode("utf-8"))

    if content_bytes > MAX_DRIVE_BYTES:
        raise ValueError(
            f"File content is too large ({content_bytes // (1024*1024)} MB). "
            "Maximum allowed size is 100 MB."
        )

    # Detect if Google sent an HTML "you need to sign in" page
    if content_bytes > 0 and "<html" in content[:500].lower() and "sign in" in content[:2000].lower():
        raise ValueError(
            "This file is not publicly shared. "
            "Please set sharing to 'Anyone with the link' in Google Drive."
        )

    if not content.strip():
        raise ValueError("The document appears to be empty or could not be read.")

    # Derive filename from Content-Disposition header or fallback
    cd = resp.headers.get("content-disposition", "")
    filename = None
    if cd:
        # RFC 5987: prefer filename*=UTF-8''... over plain filename=
        import urllib.parse as _urlparse
        star_match = re.search(r"filename\*\s*=\s*(?:UTF-8|utf-8)''(.+?)(?:;|$)", cd)
        if star_match:
            filename = _urlparse.unquote(star_match.group(1).strip())
        else:
            plain_match = re.search(r'filename\s*=\s*"?([^";]+)"?', cd)
            if plain_match:
                filename = plain_match.group(1).strip()
    if not filename:
        type_names = {
            "document": "Google Doc",
            "spreadsheet": "Google Sheet",
            "presentation": "Google Slides",
            "file": "Drive File",
        }
        filename = f"{type_names.get(used_url_type, 'Drive File')}_{file_id[:8]}.txt"

    # Determine doc_type
    if used_url_type == "spreadsheet":
        doc_type = "txt"  # CSV is plain text
    else:
        doc_type = "text"

    # Build LangChain document and split into chunks
    lc_doc = LCDocument(
        page_content=content,
        metadata={"source": url},
    )

    chunks = split_documents([lc_doc])
    if not chunks:
        raise ValueError("Document produced no chunks after splitting.")

    enriched = enrich_metadata(
        chunks=chunks,
        team_id=team_id,
        document_id=document_id,
        filename=filename,
        uploader_email=uploader_email,
        doc_type=doc_type,
    )

    return enriched, filename, doc_type, content_bytes


def validate_public_http_url(url: str) -> str:
    """
    Validate URL is HTTP(S) and resolves only to public IP addresses.
    Raises ValueError if unsafe.
    Returns normalized URL string on success.
    """
    parsed = urlparse(url.strip())
    if parsed.scheme not in {"http", "https"}:
        raise ValueError("URL must start with http:// or https://")
    if not parsed.hostname:
        raise ValueError("URL must include a valid hostname.")

    host = parsed.hostname

    def _is_public_ip(ip_str: str) -> bool:
        ip_obj = ipaddress.ip_address(ip_str)
        return not (
            ip_obj.is_private
            or ip_obj.is_loopback
            or ip_obj.is_link_local
            or ip_obj.is_multicast
            or ip_obj.is_reserved
            or ip_obj.is_unspecified
        )

    # If host is a literal IP, validate directly.
    try:
        if not _is_public_ip(host):
            raise ValueError("URL host resolves to a non-public IP address.")
        return parsed.geturl()
    except ValueError:
        # Not an IP literal, continue with DNS resolution.
        pass

    # Resolve DNS and ensure every resolved address is public.
    try:
        addrinfo = socket.getaddrinfo(host, parsed.port or (443 if parsed.scheme == "https" else 80), proto=socket.IPPROTO_TCP)
    except socket.gaierror:
        raise ValueError("Could not resolve URL hostname.")

    resolved_ips = {ai[4][0] for ai in addrinfo}
    if not resolved_ips:
        raise ValueError("Could not resolve URL hostname.")

    for ip_str in resolved_ips:
        try:
            if not _is_public_ip(ip_str):
                raise ValueError("URL host resolves to a non-public IP address.")
        except ValueError:
            raise ValueError("URL host resolves to an invalid IP address.")

    return parsed.geturl()


def process_url(
    url: str,
    team_id: int,
    document_id: int,
    uploader_email: str,
) -> Tuple[List[LCDocument], str, int]:
    """
    Fetch a generic URL page and process its text content into chunks.

    Returns: (chunks, title, file_size_bytes)

    Raises ValueError on failure.
    """
    validated_url = validate_public_http_url(url)
    try:
        with httpx.Client(follow_redirects=False, timeout=30.0) as client:
            resp = client.get(validated_url, headers={"User-Agent": "LogiPlanner/1.0"})
    except httpx.TimeoutException:
        raise ValueError("Request timed out when fetching the URL.")
    except httpx.RequestError as e:
        raise ValueError(f"Failed to fetch URL: {e}")

    if resp.status_code != 200:
        raise ValueError(f"URL returned HTTP {resp.status_code}.")

    content = resp.text
    if not content.strip():
        raise ValueError("The URL returned no content.")

    # Simple HTML text extraction — strip tags
    if "<html" in content[:500].lower() or "<body" in content[:500].lower():
        content = re.sub(r"<script[^>]*>.*?</script>", "", content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL | re.IGNORECASE)
        content = re.sub(r"<[^>]+>", " ", content)
        content = re.sub(r"\s+", " ", content).strip()

    if not content.strip():
        raise ValueError("No readable text could be extracted from the URL.")

    content_bytes = len(content.encode("utf-8"))

    # Extract title from URL or page
    title_match = re.search(r"<title[^>]*>(.*?)</title>", resp.text, re.IGNORECASE | re.DOTALL)
    title = title_match.group(1).strip() if title_match else url.split("/")[-1][:60] or "Web Page"

    lc_doc = LCDocument(
        page_content=content,
        metadata={"source": url},
    )

    chunks = split_documents([lc_doc])
    if not chunks:
        raise ValueError("Page produced no chunks after splitting.")

    enriched = enrich_metadata(
        chunks=chunks,
        team_id=team_id,
        document_id=document_id,
        filename=title,
        uploader_email=uploader_email,
        doc_type="text",
    )

    return enriched, title, content_bytes


def normalize_github_url(url: str) -> tuple[str, str]:
    """Convert a GitHub file page URL into a raw content URL and return a filename."""
    github_url = url.strip()
    if not github_url.startswith(("http://", "https://")):
        raise ValueError("GitHub URL must start with http:// or https://")

    if github_url.startswith("https://github.com/") or github_url.startswith("http://github.com/"):
        if "/blob/" not in github_url:
            raise ValueError(
                "GitHub URL must point to a file page, e.g. "
                "https://github.com/<owner>/<repo>/blob/<branch>/path/to/file"
            )
        raw_path = github_url.split("github.com/", 1)[1].replace("/blob/", "/")
        raw_url = "https://raw.githubusercontent.com/" + raw_path
        filename = os.path.basename(raw_url.split("?", 1)[0])
        if not filename:
            raise ValueError("GitHub file URL must include a filename.")
        return raw_url, filename

    if github_url.startswith("https://raw.githubusercontent.com/") or github_url.startswith("http://raw.githubusercontent.com/"):
        filename = os.path.basename(github_url.split("?", 1)[0])
        if not filename:
            raise ValueError("GitHub raw URL must include a filename.")
        return github_url, filename

    raise ValueError(
        "Only public GitHub file URLs are supported. Use a GitHub file page ending with /blob/<branch>/<file> "
        "or a raw.githubusercontent.com URL."
    )


def process_github_url(
    github_url: str,
    team_id: int,
    document_id: int,
    uploader_email: str,
) -> Tuple[List[LCDocument], str, int]:
    """Fetch GitHub file content and process it into document chunks."""
    raw_url, filename = normalize_github_url(github_url)

    try:
        with httpx.Client(follow_redirects=True, timeout=30.0) as client:
            resp = client.get(raw_url, headers={"User-Agent": "LogiPlanner/1.0"})
    except httpx.TimeoutException:
        raise ValueError("Request timed out when fetching the GitHub file.")
    except httpx.RequestError as e:
        raise ValueError(f"Failed to fetch GitHub file: {e}")

    if resp.status_code != 200:
        raise ValueError(f"GitHub file request returned HTTP {resp.status_code}.")

    content = resp.text
    if not content.strip():
        raise ValueError("GitHub file contains no readable content.")

    content_bytes = len(content.encode("utf-8"))
    doc_type = get_doc_type(filename) or "text"

    lc_doc = LCDocument(
        page_content=content,
        metadata={"source": github_url},
    )

    chunks = split_documents([lc_doc])
    if not chunks:
        raise ValueError("GitHub file produced no chunks after splitting.")

    enriched = enrich_metadata(
        chunks=chunks,
        team_id=team_id,
        document_id=document_id,
        filename=filename,
        uploader_email=uploader_email,
        doc_type=doc_type,
    )

    return enriched, filename, content_bytes


def _parse_github_owner_repo(repo_url: str) -> tuple[str, str]:
    """Extract (owner, repo) from a github.com URL."""
    repo_url = repo_url.rstrip("/")
    without_scheme = repo_url.split("github.com/", 1)
    if len(without_scheme) < 2:
        raise ValueError("URL must be a GitHub repository URL (https://github.com/owner/repo)")
    parts = without_scheme[1].split("/")
    if len(parts) < 2 or not parts[0] or not parts[1]:
        raise ValueError("URL must be a GitHub repository URL (https://github.com/owner/repo)")
    return parts[0], parts[1]


def fetch_github_repo_files(repo_url: str) -> tuple[List[Dict[str, Any]], str]:
    """Fetch repository files via the GitHub REST API using httpx.

    Returns a list of dicts with keys ``path``, ``content``, ``size`` and the
    default branch name of the repository.
    """
    owner, repo = _parse_github_owner_repo(repo_url)

    api_headers = {
        "User-Agent": "LogiPlanner/1.0",
        "Accept": "application/vnd.github+json",
    }
    if settings.GITHUB_TOKEN:
        api_headers["Authorization"] = f"Bearer {settings.GITHUB_TOKEN}"

    with httpx.Client(follow_redirects=True, timeout=30.0) as client:
        # 1. Resolve the default branch
        repo_resp = client.get(
            f"https://api.github.com/repos/{owner}/{repo}",
            headers=api_headers,
        )
        if repo_resp.status_code == 404:
            raise ValueError(f"Repository not found: {repo_url}")
        if repo_resp.status_code != 200:
            raise ValueError(
                f"GitHub API returned HTTP {repo_resp.status_code} for repository info."
            )
        default_branch = repo_resp.json().get("default_branch", "main")

        # 2. Fetch the full recursive file tree
        tree_resp = client.get(
            f"https://api.github.com/repos/{owner}/{repo}/git/trees/{default_branch}?recursive=1",
            headers=api_headers,
        )
        if tree_resp.status_code != 200:
            raise ValueError(
                f"Failed to fetch repository tree: HTTP {tree_resp.status_code}."
            )

        tree_items = tree_resp.json().get("tree", [])

        # 3. Filter to allowed blobs
        candidate_items = [
            item for item in tree_items
            if item.get("type") == "blob"
            and os.path.splitext(item["path"])[1].lower() in ALLOWED_EXTENSIONS
            and not any(part in IGNORE_DIRS for part in item["path"].split("/"))
            and item.get("size", 0) <= MAX_FILE_SIZE_REPO
        ][:MAX_REPO_FILES]

        if not candidate_items:
            raise ValueError(
                "No readable files found in the repository matching the allowed extensions."
            )

        # 4. Download each file via raw.githubusercontent.com
        collected: List[Dict[str, Any]] = []
        for item in candidate_items:
            raw_url = (
                f"https://raw.githubusercontent.com/{owner}/{repo}"
                f"/{default_branch}/{item['path']}"
            )
            try:
                content_resp = client.get(
                    raw_url,
                    headers={"User-Agent": "LogiPlanner/1.0"},
                )
                if content_resp.status_code != 200:
                    continue
                content = content_resp.text
                if content.strip():
                    collected.append({
                        "path": item["path"],
                        "content": content,
                        "size": item.get("size", len(content.encode("utf-8"))),
                    })
            except (httpx.TimeoutException, httpx.RequestError):
                continue

    return collected, default_branch


def process_github_repo(
    repo_url: str,
    team_id: int,
    document_id: int,
    uploader_email: str,
) -> Tuple[List[LCDocument], str, int]:
    """Fetch a public GitHub repo via the REST API and process all valid files into chunks."""
    # Extract repo name for the document title
    repo_url = repo_url.rstrip("/")
    repo_name = repo_url.split("/")[-1]
    if repo_name.endswith(".git"):
        repo_name = repo_name[:-4]

    files, default_branch = fetch_github_repo_files(repo_url)

    if not files:
        raise ValueError("No readable files found in the repository")

    all_chunks: List[LCDocument] = []
    total_size = 0

    for file_info in files:
        lc_doc = LCDocument(
            page_content=file_info["content"],
            metadata={"source": f"{repo_url}/blob/{default_branch}/{file_info['path']}"},
        )

        chunks = split_documents([lc_doc])
        if chunks:
            enriched = enrich_metadata(
                chunks=chunks,
                team_id=team_id,
                document_id=document_id,
                filename=file_info["path"],
                uploader_email=uploader_email,
                doc_type=get_doc_type(file_info["path"]) or "text",
            )
            all_chunks.extend(enriched)
            total_size += file_info["size"]

    if not all_chunks:
        raise ValueError("Repository files produced no chunks after splitting")

    return all_chunks, f"{repo_name} (repo)", total_size
